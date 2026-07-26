"""审查报告导出 Word：把 Agent 的 Markdown 报告转成正式格式的 .docx。

土木行业交付物以 Word 为主，这一步是"AI 输出"到"能上会的报告"的最后一公里。
只处理 Agent 报告实际会出现的 Markdown 子集：## 标题、- 列表、**加粗**、普通段落。
"""
import re
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt

BOLD_RE = re.compile(r"\*\*(.+?)\*\*")


def _set_font(run, size=12, bold=False):
    """中文字体要同时设置 east asia 名称，否则 Word 里显示成宋体默认。"""
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "仿宋_GB2312")
    run.font.size = Pt(size)
    run.font.bold = bold


def _add_markdown_line(doc, text: str):
    """把一行 Markdown 写入 docx，处理 **加粗** 片段。"""
    if text.startswith("## "):
        p = doc.add_paragraph()
        run = p.add_run(text[3:].strip())
        _set_font(run, size=15, bold=True)
        return
    if text.startswith("# "):
        p = doc.add_paragraph()
        run = p.add_run(text[2:].strip())
        _set_font(run, size=16, bold=True)
        return

    p = doc.add_paragraph()
    if text.startswith(("- ", "* ")):
        p.style = doc.styles["List Bullet"]
        text = text[2:]
    # 按 **…** 切开，交替普通/加粗
    pos = 0
    for m in BOLD_RE.finditer(text):
        if m.start() > pos:
            _set_font(p.add_run(text[pos:m.start()]))
        _set_font(p.add_run(m.group(1)), bold=True)
        pos = m.end()
    if pos < len(text):
        _set_font(p.add_run(text[pos:]))


def export_report(markdown_report: str, project_desc: str,
                  out_dir: str | Path = "reports") -> str:
    """生成正式审查报告 docx，返回文件路径。"""
    out_dir = Path(out_dir)
    out_dir.mkdir(parents=True, exist_ok=True)

    doc = Document()

    # 标题
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("施工方案规范审查报告")
    _set_font(run, size=22, bold=True)

    # 报告信息
    now = datetime.now()
    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_font(info.add_run(f"报告编号：SR-{now:%Y%m%d-%H%M}    "
                           f"生成日期：{now:%Y年%m月%d日}"), size=10)

    doc.add_paragraph()
    p = doc.add_paragraph()
    _set_font(p.add_run("一、审查对象"), size=15, bold=True)
    _set_font(doc.add_paragraph().add_run(project_desc.strip()))

    doc.add_paragraph()
    p = doc.add_paragraph()
    _set_font(p.add_run("二、审查意见"), size=15, bold=True)
    for line in markdown_report.splitlines():
        if line.strip():
            _add_markdown_line(doc, line.strip())

    doc.add_paragraph()
    note = doc.add_paragraph()
    _set_font(note.add_run(
        "说明：本报告由施工规范审查 Agent 基于已入库规范自动生成，"
        "所引条文均标注出处；重要结论请以规范原文为准，并经专业工程师复核。"), size=10)

    path = out_dir / f"审查报告_{now:%Y%m%d_%H%M%S}.docx"
    doc.save(str(path))
    return str(path)

"""文档摄取：把 PDF（文字版/扫描版）、docx 统一解析成带页码的文本。

输出统一结构：list[dict]，每页一个 dict：
    {"page": 页码(从1起), "text": 正文文本, "tables": [Markdown表格...], "is_ocr": 是否OCR页}
"""
import io
from pathlib import Path

import fitz  # PyMuPDF
import pdfplumber

from config import OCR_MIN_TEXT_CHARS, OCR_DPI

_ocr_engine = None  # 懒加载：OCR 模型初始化慢，用到扫描页才加载


def _get_ocr():
    global _ocr_engine
    if _ocr_engine is None:
        # rapidocr 3.x 和 rapidocr_onnxruntime 1.x 的 API 不同，都试一下
        try:
            from rapidocr import RapidOCR
            _ocr_engine = ("v3", RapidOCR())
        except ImportError:
            from rapidocr_onnxruntime import RapidOCR
            _ocr_engine = ("v1", RapidOCR())
    return _ocr_engine


def _ocr_image(png_bytes: bytes) -> str:
    """对一页的渲染图跑 OCR，返回拼接后的文本。"""
    version, engine = _get_ocr()
    if version == "v3":
        result = engine(png_bytes)
        txts = getattr(result, "txts", None) or []
        return "\n".join(txts)
    else:
        result, _ = engine(png_bytes)
        if not result:
            return ""
        return "\n".join(line[1] for line in result)


def _table_to_markdown(table: list[list]) -> str:
    """pdfplumber 提取的表格（二维列表）转 Markdown。"""
    if not table or not table[0]:
        return ""
    clean = [[(cell or "").replace("\n", " ").strip() for cell in row] for row in table]
    header = clean[0]
    lines = ["| " + " | ".join(header) + " |",
             "| " + " | ".join(["---"] * len(header)) + " |"]
    for row in clean[1:]:
        # 补齐/截断到表头列数，避免错位
        row = (row + [""] * len(header))[: len(header)]
        lines.append("| " + " | ".join(row) + " |")
    return "\n".join(lines)


def _load_pdf(path: Path) -> list[dict]:
    pages = []
    doc = fitz.open(path)
    try:
        # pdfplumber 只为表格服务；扫描版 PDF 提不出表格，失败就跳过
        try:
            plumber = pdfplumber.open(path)
        except Exception:
            plumber = None

        for i, page in enumerate(doc):
            text = page.get_text().strip()
            is_ocr = False

            if len(text) < OCR_MIN_TEXT_CHARS:
                # 文字太少 → 判定为扫描页，渲染成图交给 OCR
                pix = page.get_pixmap(dpi=OCR_DPI)
                text = _ocr_image(pix.tobytes("png")).strip()
                is_ocr = True

            tables = []
            if plumber and not is_ocr and i < len(plumber.pages):
                try:
                    for t in plumber.pages[i].extract_tables():
                        md = _table_to_markdown(t)
                        if md:
                            tables.append(md)
                except Exception:
                    pass  # 单页表格提取失败不影响整体

            pages.append({"page": i + 1, "text": text, "tables": tables, "is_ocr": is_ocr})

        if plumber:
            plumber.close()
    finally:
        doc.close()
    return pages


def _load_docx(path: Path) -> list[dict]:
    """docx 没有页码概念，整个文件当作一"页"。"""
    from docx import Document as DocxDocument

    doc = DocxDocument(str(path))
    texts = [p.text for p in doc.paragraphs if p.text.strip()]

    tables = []
    for t in doc.tables:
        rows = [[cell.text for cell in row.cells] for row in t.rows]
        md = _table_to_markdown(rows)
        if md:
            tables.append(md)

    return [{"page": 1, "text": "\n".join(texts), "tables": tables, "is_ocr": False}]


def load_document(path: str | Path) -> list[dict]:
    """入口：按扩展名分发。不支持的格式直接抛错，提示用户转换。"""
    path = Path(path)
    suffix = path.suffix.lower()
    if suffix == ".pdf":
        return _load_pdf(path)
    if suffix == ".docx":
        return _load_docx(path)
    if suffix == ".doc":
        raise ValueError("暂不支持 .doc 老格式，请先用 Word 另存为 .docx 再上传。")
    raise ValueError(f"不支持的格式：{suffix}（目前支持 .pdf / .docx）")
